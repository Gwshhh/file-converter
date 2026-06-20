# -*- coding: utf-8 -*-
import base64
import os
import sys
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed

import pypandoc
import pythoncom
import win32com.client
from pdf2docx import Converter as PdfConverter
from PySide6.QtCore import QThread, Qt, QUrl, Signal
from PySide6.QtGui import (
    QAction,
    QColor,
    QDesktopServices,
    QDragEnterEvent,
    QDropEvent,
    QFont,
    QIcon,
    QLinearGradient,
    QPainter,
    QPainterPath,
    QPixmap,
)
from PySide6.QtWidgets import (
    QApplication,
    QDialog,
    QFileDialog,
    QFrame,
    QGraphicsDropShadowEffect,
    QGridLayout,
    QHBoxLayout,
    QLabel,
    QMainWindow,
    QMenu,
    QMessageBox,
    QPushButton,
    QProgressBar,
    QSizePolicy,
    QToolButton,
    QVBoxLayout,
    QWidget,
)

try:
    from updater import check_for_updates, open_download_page, get_current_version
    UPDATER_AVAILABLE = True
except ImportError:
    UPDATER_AVAILABLE = False
    def get_current_version():
        return "1.0.0"

SUPPORTED_EXTENSIONS = {".md", ".markdown", ".html", ".htm", ".docx", ".pdf", ".txt"}
FILE_DIALOG_FILTER = "支持的文件 (*.md *.markdown *.html *.htm *.docx *.pdf *.txt);;所有文件 (*.*)"
OUTPUT_FORMATS = ["pdf", "docx", "html", "md", "txt"]

FORMAT_META = {
    "pdf": {"title": "PDF", "subtitle": "便携文档", "color": "#DC2626"},
    "docx": {"title": "Word", "subtitle": "可编辑文档", "color": "#2563EB"},
    "html": {"title": "HTML", "subtitle": "网页文件", "color": "#EA580C"},
    "md": {"title": "Markdown", "subtitle": "纯文本标记", "color": "#16A34A"},
    "txt": {"title": "TXT", "subtitle": "纯文本", "color": "#64748B"},
}

CSS_CONTENT = """
body {
    font-family: "Microsoft YaHei", "Segoe UI", sans-serif;
    max-width: 900px;
    margin: 40px auto;
    padding: 20px;
    line-height: 1.8;
    color: #333;
}
h1, h2, h3 { margin-top: 28px; margin-bottom: 16px; }
h1 { font-size: 2em; border-bottom: 2px solid #eee; padding-bottom: 10px; }
h2 { font-size: 1.6em; border-bottom: 1px solid #eee; padding-bottom: 8px; }
code {
    background: #f5f5f5;
    padding: 2px 6px;
    border-radius: 3px;
    font-family: "Consolas", "Monaco", monospace;
    font-size: 0.9em;
}
pre {
    background: #f8f8f8;
    padding: 16px;
    border-radius: 5px;
    overflow-x: auto;
    border: 1px solid #e0e0e0;
}
pre code { background: none; padding: 0; }
table { border-collapse: collapse; width: 100%; margin: 20px 0; }
th, td { border: 1px solid #ddd; padding: 12px; text-align: left; }
th { background: #f5f5f5; font-weight: bold; }
img { max-width: 100%; height: auto; }
blockquote { border-left: 4px solid #ddd; padding-left: 16px; color: #666; margin: 16px 0; }
"""


def create_app_icon() -> QIcon:
    size = 128
    pixmap = QPixmap(size, size)
    pixmap.fill(Qt.GlobalColor.transparent)
    painter = QPainter(pixmap)
    painter.setRenderHint(QPainter.RenderHint.Antialiasing)

    path = QPainterPath()
    path.addRoundedRect(4, 4, size - 8, size - 8, 24, 24)
    grad = QLinearGradient(0, 0, size, size)
    grad.setColorAt(0, QColor("#2563EB"))
    grad.setColorAt(1, QColor("#7C3AED"))
    painter.setBrush(grad)
    painter.setPen(Qt.PenStyle.NoPen)
    painter.drawPath(path)

    font = QFont("Segoe UI", 36, QFont.Weight.Black)
    painter.setFont(font)
    painter.setPen(QColor("#FFFFFF"))
    painter.drawText(pixmap.rect(), Qt.AlignmentFlag.AlignCenter, "F")

    painter.end()
    return QIcon(pixmap)


APP_ICON = None


def get_app_icon() -> QIcon:
    global APP_ICON
    if APP_ICON is None:
        APP_ICON = create_app_icon()
    return APP_ICON


def get_pandoc_path():
    """获取打包后的 pandoc 路径。

    若 PyInstaller 包内捆绑了 pandoc 则返回其路径；否则返回 None，让 pypandoc
    回退到系统 PATH 上的 pandoc（避免把 PYPANDOC_PANDOC 设成不存在的路径而失败）。
    """
    if hasattr(os, "_MEIPASS"):
        bundled = os.path.join(os._MEIPASS, "pandoc", "pandoc.exe")
        if os.path.exists(bundled):
            return bundled
    return None


def docx_to_pdf_word(docx_path: str, pdf_path: str):
    """使用 Microsoft Word 将 DOCX 转为 PDF（完美保留格式）"""
    pythoncom.CoInitialize()
    word = None
    doc = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(os.path.abspath(docx_path))
        doc.SaveAs(os.path.abspath(pdf_path), FileFormat=17)
    finally:
        if doc:
            doc.Close()
        if word:
            word.Quit()
        pythoncom.CoUninitialize()


def _pdf_to_docx_pdf2docx(pdf_path: str, docx_path: str):
    """使用 pdf2docx 将 PDF 转为 DOCX。"""
    cv = PdfConverter(os.path.abspath(pdf_path))
    try:
        cv.convert(os.path.abspath(docx_path))
    finally:
        cv.close()


def _unwrap_layout_tables(docx_path: str):
    """拆掉 pdf2docx 用来做版面定位的表格，只保留真正的数据表格。

    pdf2docx 为了还原版式，会把大量正文/行内代码包进单元格表格，
    并给单元格画上 1px 边框——在 Word 里就是文字周围的细/虚线方框。
    目标文档里没有这些框，正文是流式段落、只有真正的数据表才有表格。

    本函数：1) 把嵌套在单元格里的表格全部拆出来；2) 把列数 <=2 的
    顶层表格（版面定位用）拆成普通段落；3) 给剩下的真实数据表去掉
    单元格边框，统一设置目标文档同款的浅灰表格边框 (DEE2E4)。
    """
    import docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    try:
        document = docx.Document(docx_path)
    except Exception:
        return

    body = document.element.body
    P, TBL, TC = qn("w:p"), qn("w:tbl"), qn("w:tc")

    def cols_of(tbl):
        rows = tbl.findall(qn("w:tr"))
        return max((len(r.findall(TC)) for r in rows), default=0)

    def all_tables(root):
        return root.findall(".//" + TBL)

    def table_indent(tbl):
        """表格自身的左缩进 (twips)，拆表时需补到内部段落上。"""
        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            return 0
        ind = tbl_pr.find(qn("w:tblInd"))
        if ind is None:
            return 0
        val = ind.get(qn("w:w"))
        try:
            return int(round(float(val)))
        except (TypeError, ValueError):
            return 0

    def add_left_indent(paragraph, extra):
        """给段落左缩进追加 extra(twips)，保持拆表后水平位置不变。"""
        if extra <= 0:
            return
        ppr = paragraph.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            paragraph.insert(0, ppr)
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        cur = ind.get(qn("w:left")) or ind.get(qn("w:start")) or "0"
        try:
            base = int(round(float(cur)))
        except (TypeError, ValueError):
            base = 0
        ind.set(qn("w:left"), str(base + extra))

    def unwrap(tbl):
        parent = tbl.getparent()
        idx = list(parent).index(tbl)
        extra_indent = table_indent(tbl)
        blocks = []
        for row in tbl.findall(qn("w:tr")):
            for cell in row.findall(TC):
                for child in cell.iterchildren():
                    if child.tag in (P, TBL):
                        blocks.append(child)
        for b in blocks:
            if b.tag == P:
                add_left_indent(b, extra_indent)
            b.getparent().remove(b)
        for off, b in enumerate(blocks):
            parent.insert(idx + off, b)
        parent.remove(tbl)

    def set_clean_borders(tbl):
        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        for old in tbl_pr.findall(qn("w:tblBorders")):
            tbl_pr.remove(old)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement("w:" + edge)
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "6")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), "DEE2E4")
            borders.append(e)
        tbl_pr.append(borders)

    # 1) 反复拆除嵌套在单元格内的表格
    while True:
        nested = [t for t in all_tables(body) if t.getparent().tag == TC]
        if not nested:
            break
        for t in nested:
            unwrap(t)

    # 2) 拆除列数<=2 的顶层版面定位表格
    while True:
        layout = [
            t for t in all_tables(body) if t.getparent() is body and cols_of(t) <= 2
        ]
        if not layout:
            break
        for t in layout:
            unwrap(t)

    # 3) 规范剩余真实数据表的边框
    for t in [t for t in all_tables(body) if t.getparent() is body]:
        for tcb in list(t.iter(qn("w:tcBorders"))):
            tcb.getparent().remove(tcb)
        set_clean_borders(t)

    document.save(docx_path)


def _normalize_code_shading(docx_path: str):
    """规范化行内代码底纹，模仿目标文档的效果。

    pdf2docx 会用表格单元格承载行内代码并给整个单元格上底纹，导致
    灰底铺满单元格宽度形成"文本块"；还会把底纹误加到中文等非代码
    文字上。本函数把所有单元格级、以及加在非等宽字体上的底纹全部
    去掉，只给等宽代码字体（Consolas / Courier New）的文字本身加上
    紧贴文字的浅灰底纹（F5F7FA），效果与目标文档一致。
    """
    import docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    CODE_FONTS = {"Consolas", "Consola", "Courier New", "Courier"}
    GRAY = "F5F7FA"

    try:
        document = docx.Document(docx_path)
    except Exception:
        return

    def run_font(rpr):
        if rpr is None:
            return None
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            return None
        for attr in ("w:ascii", "w:hAnsi", "w:cs"):
            val = rfonts.get(qn(attr))
            if val:
                return val
        return None

    def run_text(run_el):
        return "".join(t.text or "" for t in run_el.iter(qn("w:t")))

    def set_run_shading(rpr, fill):
        for shd in rpr.findall(qn("w:shd")):
            rpr.remove(shd)
        if fill is None:
            return
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), fill)
        rpr.append(shd)

    body = document.element.body

    # 1) 去掉所有单元格级底纹（含嵌套表格、文本框内的表格）
    for tc_pr in body.iter(qn("w:tcPr")):
        for shd in tc_pr.findall(qn("w:shd")):
            tc_pr.remove(shd)

    # 2) 规范每个 run 的底纹：仅「行内代码」（嵌在正文中的等宽片段）保留
    #    紧贴文字的浅灰底；「围栏代码块」（整行以代码起头的多行代码，如
    #    Mermaid 文字版/ASCII 结构图）一律不加底纹——与目标文档一致。
    #    判定：逐段看第一个有文字的 run，若它本身就是代码字体，说明代码顶在
    #    行首，是围栏代码块行；若代码出现在正文之后，才是行内代码。
    def first_text_run_is_code(para):
        for r in para.findall(qn("w:r")):
            if run_text(r).strip():
                return run_font(r.find(qn("w:rPr"))) in CODE_FONTS
        return False

    code_runs = []
    for para in body.iter(qn("w:p")):
        fenced = first_text_run_is_code(para)
        for run_el in para.findall(qn("w:r")):
            rpr = run_el.find(qn("w:rPr"))
            is_code = run_font(rpr) in CODE_FONTS and run_text(run_el).strip()
            if is_code and not fenced:
                if rpr is None:
                    rpr = OxmlElement("w:rPr")
                    run_el.insert(0, rpr)
                set_run_shading(rpr, GRAY)
                code_runs.append(run_el)
            elif rpr is not None:
                set_run_shading(rpr, None)

    # 3) 在行内代码与中英文之间补一个不带底纹的空格，模仿目标文档
    #    中行内代码（独立文本框）四周的留白，避免代码与正文挤在一起。
    def make_space_run():
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = " "
        r.append(t)
        return r

    def sibling_text(run_el, after):
        sib = run_el.getnext() if after else run_el.getprevious()
        if sib is None or sib.tag != qn("w:r"):
            return None
        return run_text(sib)

    for run_el in code_runs:
        prev_txt = sibling_text(run_el, after=False)
        if prev_txt and not prev_txt.endswith((" ", "　")):
            run_el.addprevious(make_space_run())
        next_txt = sibling_text(run_el, after=True)
        if next_txt and not next_txt.startswith((" ", "　")):
            run_el.addnext(make_space_run())

    document.save(docx_path)


def _rebuild_lists(docx_path: str):
    """把 pdf2docx 用「微小项目符号图片 + 制表符」或「[符号格|正文格] 双列
    表格」表示的无序列表，重建为带真实项目符号字符与悬挂缩进的普通段落。

    pdf2docx 把 Markdown 列表符号还原成约 63500~76200 EMU 的正方形小图片，
    再用制表符或单独的窄单元格定位，导致：符号与正文不在同一行而错位；双列
    符号表被 _unwrap_layout_tables 拆成两段后正文铺满整行、符号变成孤立的
    "o"。本函数必须在拆表之前运行，统一改写成「○ 正文」并设置悬挂缩进。
    """
    import docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    BULLET_L0 = "●"      # ●
    BULLET_L1 = "○"      # ○
    MAX_BULLET_EMU = 100000

    try:
        document = docx.Document(docx_path)
    except Exception:
        return
    body = document.element.body

    R = qn("w:r")
    P = qn("w:p")
    TBL = qn("w:tbl")
    TR = qn("w:tr")
    TC = qn("w:tc")

    def drawing_is_bullet(drawing):
        for ext in drawing.iter(qn("wp:extent")):
            try:
                cx = int(ext.get("cx", "0"))
                cy = int(ext.get("cy", "0"))
            except (TypeError, ValueError):
                return False
            if cx and cy and abs(cx - cy) <= 4000 and max(cx, cy) <= MAX_BULLET_EMU:
                return True
        return False

    def run_has_bullet(run_el):
        return any(drawing_is_bullet(d) for d in run_el.findall(qn("w:drawing")))

    def run_is_tab_only(run_el):
        kids = [c for c in run_el if c.tag != qn("w:rPr")]
        return bool(kids) and all(c.tag == qn("w:tab") for c in kids)

    def ptext(el):
        return "".join(t.text or "" for t in el.iter(qn("w:t")))

    def left_indent(p):
        ppr = p.find(qn("w:pPr"))
        if ppr is None:
            return 0
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            return 0
        try:
            return int(ind.get(qn("w:left"), "0"))
        except (TypeError, ValueError):
            return 0

    def make_bullet_run(glyph):
        r = OxmlElement("w:r")
        rpr = OxmlElement("w:rPr")
        rf = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:eastAsia"):
            rf.set(qn(a), "Microsoft YaHei")
        rpr.append(rf)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "24")
        rpr.append(sz)
        col = OxmlElement("w:color")
        col.set(qn("w:val"), "24292E")
        rpr.append(col)
        r.append(rpr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = glyph + " "
        r.append(t)
        return r

    def set_list_indent(p, level):
        ppr = p.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            p.insert(0, ppr)
        for tabs in ppr.findall(qn("w:tabs")):
            ppr.remove(tabs)
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            ppr.append(ind)
        left = 840 if level == 0 else 1320
        ind.set(qn("w:left"), str(left))
        ind.set(qn("w:hanging"), "360")
        ind.set(qn("w:firstLine"), "0")
        ind.set(qn("w:right"), "0")

    def prepend_bullet(p, level):
        glyph = BULLET_L0 if level == 0 else BULLET_L1
        ppr = p.find(qn("w:pPr"))
        idx = (list(p).index(ppr) + 1) if ppr is not None else 0
        p.insert(idx, make_bullet_run(glyph))
        set_list_indent(p, level)

    # ---- A) 段落内联项目符号 ----
    for p in [c for c in body if c.tag == P]:
        bullet_runs = [r for r in p.findall(R) if run_has_bullet(r)]
        if not bullet_runs:
            continue
        level = 0 if left_indent(p) <= 1000 else 1
        for br in bullet_runs:
            prev = br.getprevious()
            nxt = br.getnext()
            p.remove(br)
            if prev is not None and prev.tag == R and run_is_tab_only(prev):
                p.remove(prev)
            if nxt is not None and nxt.tag == R and run_is_tab_only(nxt):
                p.remove(nxt)
        for r in list(p.findall(R)):
            if run_is_tab_only(r):
                p.remove(r)
            else:
                break
        prepend_bullet(p, level)

    # ---- B) [符号格 | 正文格] 双列表格 ----
    for tbl in [c for c in body if c.tag == TBL]:
        rows = tbl.findall(TR)
        if not rows:
            continue
        ok = True
        for tr in rows:
            tcs = tr.findall(TC)
            if len(tcs) != 2:
                ok = False
                break
            c0 = tcs[0]
            has_b = any(run_has_bullet(r) for r in c0.iter(R))
            if not has_b or ptext(c0).strip():
                ok = False
                break
        if not ok:
            continue
        new_paras = []
        for tr in rows:
            textcell = tr.findall(TC)[1]
            cell_paras = textcell.findall(qn("w:p"))
            if not cell_paras:
                continue
            src = cell_paras[0]
            prepend_bullet(src, 1)
            new_paras.append(src)
        for np_ in new_paras:
            tbl.addprevious(np_)
        tbl.getparent().remove(tbl)

    document.save(docx_path)


def _remove_orphan_bullet_drawings(docx_path: str):
    """删除 _unwrap_layout_tables 展开后留下的「仅含项目符号小图、无文字」孤立段落。

    拆表时原来的符号格（第一列只有一个小正方形图片，无文字）被提升为
    独立段落，与正文脱钩，形成浮空的"○"或"·"。直接删掉即可。
    """
    import docx
    from docx.oxml.ns import qn

    MAX_BULLET_EMU = 100000

    try:
        document = docx.Document(docx_path)
    except Exception:
        return
    body = document.element.body

    def is_tiny_square(ext):
        try:
            cx, cy = int(ext.get("cx", 0)), int(ext.get("cy", 0))
            return cx > 0 and cy > 0 and abs(cx - cy) <= 4000 and max(cx, cy) <= MAX_BULLET_EMU
        except (TypeError, ValueError):
            return False

    to_remove = []
    for p in [c for c in body if c.tag == qn("w:p")]:
        text = "".join(t.text or "" for t in p.iter(qn("w:t"))).strip()
        if text:
            continue
        extents = p.findall(".//" + qn("wp:extent"))
        if extents and all(is_tiny_square(e) for e in extents):
            to_remove.append(p)

    for p in to_remove:
        body.remove(p)

    if to_remove:
        document.save(docx_path)


def _fix_heading_styles(docx_path: str):
    """修正 pdf2docx 将正文段落过检为 Heading2/3 并对几乎所有 run 加粗的问题。

    - Heading3 → Normal，去除段落内所有 run 的显式 w:b/w:bCs
    - Heading2 → ListParagraph（保留加粗，对应参考文档子节标签格式）
    - Heading1 / Normal（含无显式 pStyle 的隐式 Normal）→ 去除 w:b/w:bCs
    """
    import docx
    from docx.oxml.ns import qn

    try:
        doc = docx.Document(docx_path)
    except Exception:
        return

    for p in doc.element.body.iter(qn("w:p")):
        pPr = p.find(qn("w:pPr"))
        pStyle = pPr.find(qn("w:pStyle")) if pPr is not None else None
        sv = pStyle.get(qn("w:val"), "") if pStyle is not None else ""

        if sv == "Heading3":
            pStyle.set(qn("w:val"), "Normal")
            sv = "Normal"
        elif sv == "Heading2":
            pStyle.set(qn("w:val"), "ListParagraph")
            continue  # 保留加粗

        if sv in ("", "Normal", "Heading1"):
            for rPr in p.findall(".//" + qn("w:rPr")):
                # 保留大字号（≥15pt = sz≥30）的粗体：标题(sz48)、节标题(sz36)、子标题(sz30)
                sz_el = rPr.find(qn("w:sz"))
                sz_val = int(sz_el.get(qn("w:val"), "0")) if sz_el is not None else 0
                if sz_val >= 30:
                    continue
                for tag in (qn("w:b"), qn("w:bCs")):
                    for el in list(rPr.findall(tag)):
                        rPr.remove(el)

    doc.save(docx_path)


def _remove_float_decorations(docx_path: str):
    """删除 pdf2docx 把 Markdown 行内代码/高亮底纹光栅化后生成的浮动灰色
    背景图片，避免它们脱离文字、错位铺在版面上遮挡正文。

    pdf2docx 会把 PDF 里那一小块浅灰代码底纹（连同里面的文字一起）整体
    截成一张 PNG，再以 behindDoc 的浮动锚点贴回去。一旦正文回流（我们清掉
    了超大右缩进让文字重新排版），这些图片不会跟着动，就变成飘在页面上的
    灰块——正是"遮挡文字 / 排版乱"的元凶。行内代码的灰底我们已用 run 级
    shading 紧贴文字重建，故这些浮图纯属冗余。

    判定规则：删除「浮动锚点(anchor) 且渲染高度 ≤ 120px(约 4 行)」的图片。
    真正的插图/示意图（本例为 图1 流程图，高 328px）远高于此阈值，予以保留；
    随文(inline)图片一律保留。
    """
    import docx
    from docx.oxml.ns import qn

    WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    MAX_DECO_EMU = 120 * 9525  # 约 120px

    try:
        document = docx.Document(docx_path)
    except Exception:
        return

    body = document.element.body
    R = qn("w:r")

    def ancestor_run(el):
        node = el.getparent()
        while node is not None and node.tag != R:
            node = node.getparent()
        return node

    runs_to_remove = []
    for drawing in body.iter(qn("w:drawing")):
        anchor = None
        for e in drawing.iter(WP + "anchor"):
            anchor = e
            break
        if anchor is None:
            continue  # 随文图片保留
        ext = None
        for e in drawing.iter(WP + "extent"):
            ext = e
            break
        cy = int(ext.get("cy", "0")) if ext is not None else 0
        if 0 < cy <= MAX_DECO_EMU:
            run = ancestor_run(drawing)
            if run is not None:
                runs_to_remove.append(run)

    removed = 0
    seen = set()
    for run in runs_to_remove:
        if id(run) in seen:
            continue
        seen.add(id(run))
        parent = run.getparent()
        if parent is not None:
            parent.remove(run)
            removed += 1

    if removed:
        document.save(docx_path)


def _inline_vector_figures(docx_path: str, pdf_path: str):
    """把 pdf2docx「光栅化丢字 + 浮动锚点」的矢量插图，整体替换成一张
    从源 PDF 同区域高清重渲染（文字/连线/方框俱全）的「随文」图片。

    pdf2docx 处理 Mermaid 流程图这类矢量图时有两个毛病：
    (1) 只把矢量路径（方框 + 连线）截成 PNG，节点文字当独立文本层抽出去，
        于是位图里只剩空方框，文字另散落成段落；
    (2) 用 behindDoc 浮动锚点按 PDF 页面坐标定位，正文一回流图就钉在原处
        盖住表格/正文。

    关键发现：浮动锚点位置精确编码了该图在源 PDF 中的区域——positionH/V 的
    posOffset(EMU) 即裁剪框左上角、extent 即宽高（EMU÷12700=pt）；锚点之前
    出现的 sectPr 个数恰好等于源 PDF 页码（pdf2docx 每页一个分节）。据此用
    PyMuPDF 高 DPI 重渲染整块区域（含文字），再用 python-docx 原生接口把这张
    干净位图作为 inline 图片插回锚点所在段落，并删除原浮动 drawing。直接复用
    pdf2docx 的 drawing XML 改 inline 会被 Word 压成一条细缝，故必须重建。

    必须在 _normalize_sections 之前运行（此时每页一个 section，页码可由
    sectPr 计数推出）。
    """
    import io
    import docx
    from docx.oxml.ns import qn
    from docx.shared import Emu

    WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    MIN_FIG_EMU = 120 * 9525
    EMU_PER_PT = 12700.0

    try:
        import fitz
    except Exception:
        return
    try:
        document = docx.Document(docx_path)
    except Exception:
        return

    body = document.element.body
    R = qn("w:r")
    P = qn("w:p")

    def containing(el, tag):
        node = el.getparent()
        while node is not None and node.tag != tag:
            node = node.getparent()
        return node

    # 收集大尺寸浮动锚点图：页码、裁剪框、所属段落、宿主 run
    jobs = []
    seen_sect = 0
    for el in body.iter():
        if el.tag == qn("w:sectPr"):
            seen_sect += 1
            continue
        if el.tag != WP + "anchor":
            continue
        ext = el.find(WP + "extent")
        if ext is None or int(ext.get("cy", "0")) <= MIN_FIG_EMU:
            continue
        cx = int(ext.get("cx", "0"))
        cy = int(ext.get("cy", "0"))

        def offset(tag):
            pe = el.find(WP + tag)
            if pe is None or pe.get("relativeFrom") != "page":
                return None
            off = pe.find(WP + "posOffset")
            try:
                return int(off.text)
            except (TypeError, ValueError, AttributeError):
                return None

        hx, vy = offset("positionH"), offset("positionV")
        if hx is None or vy is None:
            continue
        para = containing(el, P)
        run = containing(el, R)
        if para is None:
            continue
        jobs.append({
            "page": seen_sect,
            "cx": cx, "cy": cy,
            "clip": (hx / EMU_PER_PT, vy / EMU_PER_PT,
                     (hx + cx) / EMU_PER_PT, (vy + cy) / EMU_PER_PT),
            "para": para, "run": run, "drawing": containing(el, qn("w:drawing")),
        })

    if not jobs:
        return

    pdf = fitz.open(os.path.abspath(pdf_path))
    changed = False
    try:
        for job in jobs:
            if job["page"] >= pdf.page_count:
                continue
            page = pdf[job["page"]]
            x0, y0, x1, y1 = job["clip"]
            clip = fitz.Rect(x0 - 2, y0 - 3, x1 + 2, y1 + 3) & page.rect
            pix = page.get_pixmap(matrix=fitz.Matrix(4, 4), clip=clip)
            stream = io.BytesIO(pix.tobytes("png"))

            # 删除原浮动 drawing（连同空宿主 run）
            drawing = job["drawing"]
            if drawing is not None:
                dp = drawing.getparent()
                if dp is not None:
                    dp.remove(drawing)
            run = job["run"]
            if run is not None and run.getparent() is not None:
                # 仅当该 run 已无可见内容时删除
                if not run.findall(qn("w:t")) and not run.findall(qn("w:drawing")):
                    run.getparent().remove(run)

            # 用原生接口把干净位图作为 inline 插回该段落。
            # Paragraph/Run 的 add_picture 经 _parent.part 取文档 part 建立图片关系，
            # 故用一个仅暴露 document.part 的代理对象做父级。
            class _PartShim:
                part = document.part

            target_p = job["para"]
            # pdf2docx 给该段落写了固定行高 (w:lineRule="exact" w:line=220)，
            # 会把随文大图裁成只剩一条细缝——必须改成自动行高，让行盒撑开适配图片。
            ppr = target_p.find(qn("w:pPr"))
            if ppr is not None:
                spacing = ppr.find(qn("w:spacing"))
                if spacing is not None:
                    rule = spacing.get(qn("w:lineRule"))
                    if rule == "exact" or spacing.get(qn("w:line")):
                        for a in (qn("w:line"), qn("w:lineRule")):
                            if spacing.get(a) is not None:
                                del spacing.attrib[a]

            para = docx.text.paragraph.Paragraph(target_p, _PartShim())
            new_run = para.add_run()
            new_run.add_picture(stream, width=Emu(job["cx"]))
            changed = True
    finally:
        pdf.close()

    if changed:
        document.save(docx_path)


def _remove_flowchart_text_debris(docx_path: str):
    """删除 pdf2docx 在流程图大图旁「重复抽出」的节点文字段落与定位表格。

    pdf2docx 把流程图既光栅化成一张图、又把图里每个节点的文字单独抽成段落/
    表格贴在同一版面区域，于是流程图图片下面又冒出一串重复的
    "real0_tm.dat / 观测响应重排结果 / model2_tm.dat ..." 文字和一个
    "1.DAT / 1_copy.DAT" 定位表——目标文档里没有这些重复内容（都已烘进图里）。

    定位方式：用两个稳定地标夹出流程图区域——上界是那张大 inline 插图所在段落，
    下界是 Mermaid 文字版回退说明段落（含 "Mermaid"）。删掉两者之间的所有
    含文字段落与表格，保留空段落作为间距。只有当两个地标都存在、且图在前时
    才执行，避免误删。
    """
    import docx
    from docx.oxml.ns import qn

    WP = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
    MIN_FIG_EMU = 120 * 9525

    try:
        document = docx.Document(docx_path)
    except Exception:
        return

    body = document.element.body
    kids = list(body)

    def ptext(el):
        return "".join(t.text or "" for t in el.iter(qn("w:t"))).strip()

    fig_idx = None
    for i, el in enumerate(kids):
        if el.tag != qn("w:p"):
            continue
        for inline in el.findall(".//" + WP + "inline"):
            ext = inline.find(WP + "extent")
            if ext is not None and int(ext.get("cy", "0")) > MIN_FIG_EMU:
                fig_idx = i
                break
        if fig_idx is not None:
            break

    if fig_idx is None:
        return

    mermaid_idx = None
    for i in range(fig_idx + 1, len(kids)):
        el = kids[i]
        if el.tag == qn("w:p") and "Mermaid" in ptext(el):
            mermaid_idx = i
            break

    if mermaid_idx is None or mermaid_idx <= fig_idx + 1:
        return

    removed = False
    for i in range(fig_idx + 1, mermaid_idx):
        el = kids[i]
        if el.tag == qn("w:tbl"):
            body.remove(el)
            removed = True
        elif el.tag == qn("w:p") and ptext(el):
            body.remove(el)
            removed = True

    if removed:
        document.save(docx_path)


def _strip_runaway_indents(docx_path: str):
    """清除 pdf2docx 给段落塞进去的超大右缩进（w:right），避免正文被挤成
    极窄的竖排碎块。

    pdf2docx 为了还原 PDF 里文字块的右边界，会给段落写入巨大的 w:right
    （实测高达 8640 twips ≈ 6 英寸）。在 Word 里这等于把可用行宽压到只剩
    几个字符，整段文字就一个字一个字竖着换行——正是"行内代码被拆成
    rea/l0_/tm./dat 竖排"的根因。参考文档所有段落的右缩进都是 0，
    因此这里把正文里偏大的右缩进一律清零，恢复整行回流。
    """
    import docx
    from docx.oxml.ns import qn

    try:
        document = docx.Document(docx_path)
    except Exception:
        return

    body = document.element.body
    changed = False
    for para in body.iter(qn("w:p")):
        ppr = para.find(qn("w:pPr"))
        if ppr is None:
            continue
        ind = ppr.find(qn("w:ind"))
        if ind is None:
            continue
        right = ind.get(qn("w:right")) or ind.get(qn("w:end"))
        if right is None:
            continue
        try:
            rv = int(round(float(right)))
        except (TypeError, ValueError):
            continue
        # 参考文档右缩进恒为 0；这里把任何 >0 的右缩进清零，
        # 既消除竖排碎块，又与目标排版一致。
        if rv > 0:
            ind.set(qn("w:right"), "0")
            if ind.get(qn("w:end")) is not None:
                ind.set(qn("w:end"), "0")
            changed = True

    if changed:
        document.save(docx_path)


def _normalize_sections(docx_path: str):
    """统一所有 section 的页边距，并把 pdf2docx 按 PDF 页面硬塞的
    分页 section（type=nextPage）改成连续 section（type=continuous）。

    pdf2docx 会按 PDF 的每一页生成一个 sectPr，type 缺省即 nextPage，
    强制硬分页，并把每段页边距设成各不相同、且上/下常为 0 的值——
    在 Word 里文字就贴着纸边、版面随页跳动、还比目标文档多出两页。

    目标文档所有 section 的边距是统一的 L=R=708, T=1000, B=280,
    header=footer=720, gutter=0（twips）。本函数把每个 sectPr 的 pgMar
    都改成这一套，并把除最后一个（文档级）以外的 section 全部改成
    continuous，让 Word 按内容自然回流，不再被旧的 PDF 页面边界钉死。
    """
    import docx
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    REF_MAR = {
        "left": "708", "right": "708", "top": "1000", "bottom": "280",
        "header": "720", "footer": "720", "gutter": "0",
    }

    try:
        document = docx.Document(docx_path)
    except Exception:
        return

    body = document.element.body
    sect_prs = list(body.iter(qn("w:sectPr")))
    if not sect_prs:
        return

    for idx, sectPr in enumerate(sect_prs):
        # 1) 统一页边距
        pgMar = sectPr.find(qn("w:pgMar"))
        if pgMar is None:
            pgMar = OxmlElement("w:pgMar")
            sectPr.append(pgMar)
        for k, v in REF_MAR.items():
            pgMar.set(qn("w:" + k), v)

        # 2) 除文档级（最后一个）section 外，全部改为连续，取消硬分页
        is_last = idx == len(sect_prs) - 1
        type_el = sectPr.find(qn("w:type"))
        if is_last:
            if type_el is not None:
                sectPr.remove(type_el)
        else:
            if type_el is None:
                type_el = OxmlElement("w:type")
                # w:type 须排在 pgSz 之前，安全起见插到 sectPr 开头
                sectPr.insert(0, type_el)
            type_el.set(qn("w:val"), "continuous")

    document.save(docx_path)


def _md_to_docx_highfidelity(md_path: str, pdf_path: str, docx_path: str):
    """md→Word 高保真路线：pandoc 还原原生结构，再补样式与流程图。

    pandoc 直接把 Markdown 转 docx，能拿到原生 Heading 样式、原生项目符号/
    编号列表、真正的表格与行内代码字符样式——这正是目标文档的结构地基。之后
    本函数再做两件事贴近目标外观：(1) 标题改成近黑加粗、行内代码(VerbatimChar)
    加浅灰底；(2) 把 pandoc 当普通代码块输出的 ```mermaid 段落，替换成从源 PDF
    抠出的「已渲染流程图」图片。"""
    import docx

    pandoc_path = get_pandoc_path()
    if pandoc_path:
        os.environ["PYPANDOC_PANDOC"] = pandoc_path
    pypandoc.convert_file(
        md_path, "docx", outputfile=docx_path, format="markdown",
        extra_args=["--standalone"],
    )

    document = docx.Document(docx_path)
    _md_restyle_headings_and_code(document)
    _md_bold_headings(document)
    _md_set_bullet_glyphs(document)
    _md_style_tables(document)
    _md_replace_mermaid_with_images(document, md_path, pdf_path)
    document.save(docx_path)


def _md_bold_headings(document):
    """直接给每个标题段落的所有 run 加粗(w:b + w:bCs)并设近黑色与中文字体，
    确保中文部分也呈现加粗——pandoc 标题中文 run 只继承样式，靠样式加粗常被
    字体替换吞掉，目标文档则是 run 级直接加粗。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    HEAD_COLOR = "24292E"
    body = document.element.body
    for p in body.iter(qn("w:p")):
        ppr = p.find(qn("w:pPr"))
        ps = ppr.find(qn("w:pStyle")) if ppr is not None else None
        sv = ps.get(qn("w:val")) if ps is not None else ""
        if not (sv.startswith("Heading") or sv in ("Title", "Subtitle")):
            continue
        for r in p.findall(qn("w:r")):
            rpr = r.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                r.insert(0, rpr)
            # 设置带真实粗体字形的中文字体，避免合成粗体被字体替换吞掉
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.insert(0, rfonts)
            rfonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            if rfonts.get(qn("w:hint")):
                del rfonts.attrib[qn("w:hint")]
            for tag in ("w:b", "w:bCs"):
                if rpr.find(qn(tag)) is None:
                    rpr.append(OxmlElement(tag))
            if rpr.find(qn("w:color")) is None:
                col = OxmlElement("w:color")
                col.set(qn("w:val"), HEAD_COLOR)
                rpr.append(col)


def _md_style_tables(document):
    """给所有表格加目标文档同款的细灰边框(DEE2E4)，并去掉单元格底纹。
    pandoc 默认表格只有上下两条粗线、无竖线，与目标文档的「全网格细灰线」不符。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    GRID = "DEE2E4"
    body = document.element.body
    for tbl in body.iter(qn("w:tbl")):
        tblPr = tbl.find(qn("w:tblPr"))
        if tblPr is None:
            tblPr = OxmlElement("w:tblPr")
            tbl.insert(0, tblPr)
        for old in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(old)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            e = OxmlElement("w:" + edge)
            e.set(qn("w:val"), "single")
            e.set(qn("w:sz"), "6")
            e.set(qn("w:space"), "0")
            e.set(qn("w:color"), GRID)
            borders.append(e)
        tblPr.append(borders)
        # 去掉残留单元格底纹（保持与目标文档一致的纯白单元格）
        for tcPr in tbl.iter(qn("w:tcPr")):
            for shd in tcPr.findall(qn("w:shd")):
                tcPr.remove(shd)


def _md_restyle_headings_and_code(document):
    """把 pandoc 默认样式改成贴近目标文档：标题近黑加粗(Segoe UI Semibold)、
    行内代码(VerbatimChar)加浅灰底纹。直接改样式定义，全局一次生效。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    HEAD_COLOR = "24292E"   # 目标文档标题色（近黑）
    CODE_FILL = "F5F7FA"    # 行内代码浅灰底

    def style_rpr(style):
        rpr = style.element.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            style.element.append(rpr)
        return rpr

    def set_color(rpr, val):
        for c in rpr.findall(qn("w:color")):
            rpr.remove(c)
        c = OxmlElement("w:color")
        c.set(qn("w:val"), val)
        rpr.append(c)

    def set_bold(rpr):
        if rpr.find(qn("w:b")) is None:
            rpr.append(OxmlElement("w:b"))

    def style_ppr(style):
        ppr = style.element.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            # pPr 须排在 rPr 之前
            rpr = style.element.find(qn("w:rPr"))
            if rpr is not None:
                rpr.addprevious(ppr)
            else:
                style.element.append(ppr)
        return ppr

    def set_bottom_border(ppr):
        for b in ppr.findall(qn("w:pBdr")):
            ppr.remove(b)
        pbdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")        # 0.5pt 细线
        bottom.set(qn("w:space"), "4")
        bottom.set(qn("w:color"), "EAECEF")  # GitHub 风格浅灰
        pbdr.append(bottom)
        # pBdr 须在 spacing/ind 之前，放到 pPr 开头
        ppr.insert(0, pbdr)

    styles = document.styles
    for name in ("Title", "Heading 1", "Heading 2", "Heading 3",
                 "Heading 4", "Heading 5", "Heading 6"):
        try:
            st = styles[name]
        except KeyError:
            continue
        rpr = style_rpr(st)
        set_color(rpr, HEAD_COLOR)
        set_bold(rpr)
        # 目标文档标题/一二级小节下有细灰线；正文小标题(H3+)不加，避免过密
        if name in ("Title", "Heading 1", "Heading 2"):
            set_bottom_border(style_ppr(st))

    # 行内代码：VerbatimChar 字符样式加浅灰底
    try:
        vc = styles["Verbatim Char"]
    except KeyError:
        vc = None
    if vc is None:
        for st in styles:
            if st.name and st.name.replace(" ", "").lower() == "verbatimchar":
                vc = st
                break
    if vc is not None:
        rpr = style_rpr(vc)
        for s in rpr.findall(qn("w:shd")):
            rpr.remove(s)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), CODE_FILL)
        rpr.append(shd)
        # 目标文档行内代码：Consolas、字号 10pt(sz=20，略小于正文)、紧排(spacing=-2)
        for tag in ("w:sz", "w:szCs", "w:spacing", "w:rFonts", "w:color"):
            for e in rpr.findall(qn(tag)):
                rpr.remove(e)
        rfonts = OxmlElement("w:rFonts")
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            rfonts.set(qn(a), "Consolas")
        rpr.insert(0, rfonts)
        col = OxmlElement("w:color")
        col.set(qn("w:val"), HEAD_COLOR)
        rpr.append(col)
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:val"), "-2")
        rpr.append(sp)
        for t in ("w:sz", "w:szCs"):
            e = OxmlElement(t)
            e.set(qn("w:val"), "20")
            rpr.append(e)

    # 列表项：pandoc 用 Compact 样式且段距为 0，显得拥挤；目标文档每个列表项
    # 上方有约 11pt 间距。给 Compact / List Paragraph 段前加间距、放宽行距。
    for sname in ("Compact", "List Paragraph", "Body Text", "First Paragraph"):
        try:
            st = styles[sname]
        except KeyError:
            continue
        ppr = st.element.find(qn("w:pPr"))
        if ppr is None:
            ppr = OxmlElement("w:pPr")
            rpr0 = st.element.find(qn("w:rPr"))
            if rpr0 is not None:
                rpr0.addprevious(ppr)
            else:
                st.element.append(ppr)
        for sp in ppr.findall(qn("w:spacing")):
            ppr.remove(sp)
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "120")
        sp.set(qn("w:after"), "120")
        sp.set(qn("w:line"), "288")          # 1.2x 行距
        sp.set(qn("w:lineRule"), "auto")
        ppr.append(sp)


def _md_set_bullet_glyphs(document):
    """把无序列表的项目符号字形改成目标文档同款：0 级实心菱形 ◆、1 级空心圆 ○、
    2 级实心方块 ▪，并调小字号、改灰色（目标文档符号很小、偏灰），避免又大又黑
    （pandoc 默认是 ●/o/▪ 全黑大号）。"""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    GLYPHS = {"0": "◆", "1": "○", "2": "▪"}
    GLYPH_COLOR = "586069"   # 目标文档符号偏灰
    GLYPH_SZ = "12"          # 6pt，贴近目标文档很小的菱形
    INDENT = {"0": ("420", "280"), "1": ("840", "280"), "2": ("1260", "280")}
    try:
        numbering = document.part.numbering_part.element
    except Exception:
        return
    for lvl in numbering.iter(qn("w:lvl")):
        ilvl = lvl.get(qn("w:ilvl"))
        fmt = lvl.find(qn("w:numFmt"))
        if fmt is None or fmt.get(qn("w:val")) != "bullet":
            continue
        glyph = GLYPHS.get(ilvl)
        if glyph is None:
            continue
        lt = lvl.find(qn("w:lvlText"))
        if lt is not None:
            lt.set(qn("w:val"), glyph)
        # 收紧缩进，让符号靠近左边距（目标文档符号紧贴正文左侧）
        if ilvl in INDENT:
            ppr = lvl.find(qn("w:pPr"))
            if ppr is None:
                ppr = OxmlElement("w:pPr")
                lvl.append(ppr)
            ind = ppr.find(qn("w:ind"))
            if ind is None:
                ind = OxmlElement("w:ind")
                ppr.append(ind)
            left, hang = INDENT[ilvl]
            ind.set(qn("w:left"), left)
            ind.set(qn("w:hanging"), hang)
        rpr = lvl.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            lvl.append(rpr)
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            rpr.insert(0, rfonts)
        for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(a), "MS Gothic")
        if rfonts.get(qn("w:hint")):
            del rfonts.attrib[qn("w:hint")]
        for tag in ("w:color", "w:sz", "w:szCs"):
            for e in rpr.findall(qn(tag)):
                rpr.remove(e)
        col = OxmlElement("w:color")
        col.set(qn("w:val"), GLYPH_COLOR)
        rpr.append(col)
        for t in ("w:sz", "w:szCs"):
            e = OxmlElement(t)
            e.set(qn("w:val"), GLYPH_SZ)
            rpr.append(e)


def _md_replace_mermaid_with_images(document, md_path, pdf_path):
    """把 pandoc 当代码块输出的 Mermaid 段落，替换成从源 PDF 抠出的流程图图片。

    pandoc 把 ```mermaid 围栏整段放进一个 SourceCode 段落（文内用换行分隔）。
    本函数按出现顺序找出这些「以 mermaid 图型关键字开头」的 SourceCode 段落，
    依次用 _locate_mermaid_images 返回的 PNG 替换；定位失败的保持原样（仍是
    可读的文字版）。普通代码块（ASCII 结构图、数据样例）不受影响。"""
    from docx.oxml.ns import qn
    from docx.shared import Emu

    MERMAID_KW = ("flowchart", "graph ", "graph\t", "sequencediagram",
                  "classdiagram", "statediagram", "erdiagram", "gantt",
                  "pie", "journey", "gitgraph", "mindmap", "timeline")

    body = document.element.body

    def para_text(p):
        return "".join(t.text or "" for t in p.iter(qn("w:t")))

    def is_mermaid(p):
        ppr = p.find(qn("w:pPr"))
        ps = ppr.find(qn("w:pStyle")) if ppr is not None else None
        sv = ps.get(qn("w:val")) if ps is not None else ""
        if sv != "SourceCode":
            return False
        head = para_text(p).lstrip().lower()
        return any(head.startswith(k) for k in MERMAID_KW)

    targets = [p for p in body.iter(qn("w:p")) if is_mermaid(p)]
    if not targets:
        return

    try:
        images = _locate_mermaid_images(md_path, pdf_path)
    except Exception:
        images = []

    # 目标页宽（twips→EMU）：用文档第一个 section 的可用正文宽
    try:
        sect = document.sections[0]
        avail_emu = int(sect.page_width) - int(sect.left_margin) - int(sect.right_margin)
    except Exception:
        avail_emu = 5800000

    import io
    from PIL import Image

    for idx, p in enumerate(targets):
        png = images[idx] if idx < len(images) else None
        if not png:
            continue
        # 清空原代码文字 run
        for r in list(p.findall(qn("w:r"))):
            p.remove(r)
        # 段落改为居中、普通行高，避免 SourceCode 的等宽底纹/固定行高影响图片
        ppr = p.find(qn("w:pPr"))
        if ppr is not None:
            ps = ppr.find(qn("w:pStyle"))
            if ps is not None:
                ppr.remove(ps)
            for sp in ppr.findall(qn("w:spacing")):
                for a in (qn("w:line"), qn("w:lineRule")):
                    if sp.get(a) is not None:
                        del sp.attrib[a]
        # 按宽度等比缩放（不超过正文可用宽）
        try:
            w_px, h_px = Image.open(io.BytesIO(png)).size
        except Exception:
            w_px, h_px = 1000, 500
        target_w = min(avail_emu, int(w_px * 9525))  # 9525 EMU/px @96dpi
        para = _docx_paragraph(document, p)
        run = para.add_run()
        run.add_picture(io.BytesIO(png), width=Emu(target_w))


def _docx_paragraph(document, p_element):
    """用既有 <w:p> 元素构造 python-docx Paragraph（父级暴露 document.part）。"""
    import docx

    class _Shim:
        part = document.part

    return docx.text.paragraph.Paragraph(p_element, _Shim())


_MD_SKIP_DIRS = {
    ".venv", "venv", "node_modules", ".git", "__pycache__",
    "site-packages", "dist", "build", ".cache", ".idea", ".vscode",
}


def _find_markdown_source(pdf_path: str, max_depth: int = 7, cap: int = 20000):
    """在 PDF 所在目录树里就近查找「同名 Markdown 源文件」。

    很多 PDF 是由 Markdown 渲染而来（如本项目的 WFEM 说明）。若能找到对应
    的 .md 源，就能用「md→Word」拿到原生标题/列表/编号结构，远比从 PDF 反推
    版面更贴近真正的目标文档。查找以 PDF 同目录为根、按文件名主干（stem）匹配，
    限定深度并跳过依赖/缓存等重目录，命中即返回；找不到返回 None（回退到
    pdf2docx 路线）。还会做一次内容校验：md 里需出现 PDF 文本的若干特征行，
    避免误配同名但无关的文件。"""
    stem = os.path.splitext(os.path.basename(pdf_path))[0]
    root = os.path.dirname(os.path.abspath(pdf_path)) or "."
    root_depth = root.rstrip("/\\").count(os.sep)
    candidates = []
    n = 0
    for dp, dns, fns in os.walk(root):
        dns[:] = [d for d in dns if d not in _MD_SKIP_DIRS and not d.startswith(".")]
        if dp.count(os.sep) - root_depth > max_depth:
            dns[:] = []
            continue
        n += 1
        if n > cap:
            break
        if stem + ".md" in fns:
            candidates.append(os.path.join(dp, stem + ".md"))
    if not candidates:
        return None
    # 内容校验：取 PDF 前两页若干中文/代码特征词，确认 md 确为该 PDF 的源
    try:
        import fitz
        doc = fitz.open(os.path.abspath(pdf_path))
        sample = "".join(doc[i].get_text() for i in range(min(2, doc.page_count)))
        doc.close()
    except Exception:
        sample = ""
    keys = [w for w in __import__("re").findall(r"[一-鿿]{4,}", sample)][:8]
    best = candidates[0]
    if keys:
        best_hits = -1
        for c in candidates:
            try:
                txt = open(c, encoding="utf-8", errors="ignore").read()
            except OSError:
                continue
            hits = sum(1 for k in keys if k in txt)
            if hits > best_hits:
                best_hits, best = hits, c
        if best_hits <= 0:
            return None
    return best


def _locate_mermaid_images(md_path: str, pdf_path: str):
    """用 md 里每个 ```mermaid 围栏「前后相邻文字」作锚点，在源 PDF 中定位
    该流程图被渲染出来的版面区域，高 DPI 裁成 PNG，按出现顺序返回字节列表。

    Markdown→PDF 时 Mermaid 被渲染成矢量图；pandoc 却把 ```mermaid 当普通代码
    块输出成文字。于是这里直接从源 PDF 把「画好的图」抠出来：流程图夹在它前
    一段标题/正文与后一段说明文字之间，定位这两段文字在 PDF 的位置即可框出
    图所在的竖直区间，再取该区间内所有图形/文字的并集 bbox 精确裁剪。"""
    import re
    import fitz

    try:
        md = open(md_path, encoding="utf-8", errors="ignore").read()
    except OSError:
        return []
    lines = md.splitlines()

    def clean(s):
        return re.sub(r"[#*`>|~\-]", "", s).strip()

    blocks = []
    i = 0
    while i < len(lines):
        if lines[i].strip().startswith("```mermaid"):
            j = i - 1
            while j >= 0 and not lines[j].strip():
                j -= 1
            before = clean(lines[j]) if j >= 0 else ""
            k = i + 1
            while k < len(lines) and not lines[k].strip().startswith("```"):
                k += 1
            m = k + 1
            while m < len(lines) and not lines[m].strip():
                m += 1
            after = clean(lines[m]) if m < len(lines) else ""
            blocks.append((before, after))
            i = k + 1
        else:
            i += 1
    if not blocks:
        return []

    doc = fitz.open(os.path.abspath(pdf_path))

    def search(needle):
        key = needle[:12]
        hits = []
        if not key:
            return hits
        for pno in range(doc.page_count):
            for r in doc[pno].search_for(key):
                hits.append((pno, r))
        return hits

    images = []
    for before, after in blocks:
        bh = search(before)
        ah = search(after)
        if not ah:
            images.append(None)
            continue
        ap, ar = ah[0]
        # before 锚点：取与 after 同页或其前一页、且在 after 之上的最近一个
        top_y = None
        same = [r for (p, r) in bh if p == ap and r.y1 <= ar.y0]
        if same:
            top_y = max(r.y1 for r in same)
        page = doc[ap]
        gap = 8
        x0 = y0 = 1e9
        x1 = y1 = -1e9
        lo = (top_y + gap) if top_y is not None else 0
        hi = ar.y0 - gap
        for d in page.get_drawings():
            r = d["rect"]
            if r.y0 >= lo and r.y1 <= hi and r.width < 560 and r.height < 420:
                x0, y0 = min(x0, r.x0), min(y0, r.y0)
                x1, y1 = max(x1, r.x1), max(y1, r.y1)
        for b in page.get_text("blocks"):
            bx0, by0, bx1, by1 = b[:4]
            if by0 >= lo and by1 <= hi:
                x0, y0 = min(x0, bx0), min(y0, by0)
                x1, y1 = max(x1, bx1), max(y1, by1)
        if x1 <= x0 or y1 <= y0:
            images.append(None)
            continue
        clip = fitz.Rect(x0 - 3, y0 - 3, x1 + 3, y1 + 3) & page.rect
        pix = page.get_pixmap(matrix=fitz.Matrix(4, 4), clip=clip)
        images.append(pix.tobytes("png"))
    doc.close()
    return images


def pdf_to_docx_word(pdf_path: str, docx_path: str):
    """将 PDF 转为 DOCX（智能选择转换路线）。

    若能在 PDF 同目录树就近找到「同名 Markdown 源文件」，走 md→Word 高保真
    路线（pandoc 还原原生标题/列表/编号结构，再补样式与流程图），最贴近真正
    的目标文档；否则回退到 pdf2docx 路线（不依赖 md 源，任意 PDF 都能转）。"""
    md_src = None
    try:
        md_src = _find_markdown_source(pdf_path)
    except Exception:
        md_src = None
    if md_src:
        try:
            _md_to_docx_highfidelity(md_src, pdf_path, docx_path)
            return
        except Exception:
            pass  # 高保真路线异常时回退到 pdf2docx
    _pdf_to_docx_from_pdf(pdf_path, docx_path)


def _pdf_to_docx_from_pdf(pdf_path: str, docx_path: str):
    """从 PDF 反推版面的转换路线（无 md 源时使用）。

    使用 pdf2docx 转换（保留流式结构，不会像 Word 原生重排那样把
    文字摆成相互重叠的浮动文本框），先把项目符号图片重建为真实列表
    段落（必须在拆表前执行），再拆除版面定位表格（顺带清除拆表后
    遗留的孤立符号段落），规范化行内代码底纹，最后修正过度检测的
    标题样式与错误加粗，最后统一各 section 的页边距并把按 PDF 页面
    硬分页的 section 改为连续，让版面按内容自然回流、不再贴边或多页。
    """
    _pdf_to_docx_pdf2docx(pdf_path, docx_path)
    _inline_vector_figures(docx_path, pdf_path)
    _remove_float_decorations(docx_path)
    _remove_flowchart_text_debris(docx_path)
    _rebuild_lists(docx_path)
    _unwrap_layout_tables(docx_path)
    _remove_orphan_bullet_drawings(docx_path)
    _normalize_code_shading(docx_path)
    _fix_heading_styles(docx_path)
    _strip_runaway_indents(docx_path)
    _normalize_sections(docx_path)


def convert_with_pandoc(input_path: str, output_format: str, output_path: str):
    """使用 Pandoc 进行格式转换"""
    ext = os.path.splitext(input_path)[1].lower()

    if ext == ".docx" and output_format == "pdf":
        docx_to_pdf_word(input_path, output_path)
        return

    if ext == ".pdf" and output_format == "docx":
        pdf_to_docx_word(input_path, output_path)
        return

    if ext == ".pdf" and output_format == "pdf":
        import shutil

        shutil.copy2(input_path, output_path)
        return

    format_map = {
        ".md": "markdown",
        ".markdown": "markdown",
        ".html": "html",
        ".htm": "html",
        ".docx": "docx",
        ".txt": "plain",
        ".pdf": None,
    }
    input_format = format_map.get(ext, "markdown")

    if ext == ".pdf":
        docx_tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".docx", delete=False)
        docx_tmp.close()
        pdf_to_docx_word(input_path, docx_tmp.name)
        input_path = docx_tmp.name
        input_format = "docx"

    output_format_map = {
        "pdf": "docx",
        "docx": "docx",
        "html": "html",
        "md": "markdown",
        "txt": "plain",
    }
    pandoc_output_format = output_format_map.get(output_format, "html")

    resource_path = os.path.dirname(os.path.abspath(input_path))
    extra_args = [
        "--standalone",
        "--embed-resources",
        f"--resource-path={resource_path}",
    ]

    pandoc_path = get_pandoc_path()
    if pandoc_path:
        os.environ["PYPANDOC_PANDOC"] = pandoc_path

    if output_format == "pdf":
        docx_tmp = tempfile.NamedTemporaryFile(mode="w", suffix=".docx", delete=False)
        docx_tmp.close()
        pypandoc.convert_file(
            input_path,
            "docx",
            outputfile=docx_tmp.name,
            format=input_format,
            extra_args=extra_args,
        )
        docx_to_pdf_word(docx_tmp.name, output_path)
        try:
            os.unlink(docx_tmp.name)
        except OSError:
            pass
    else:
        if output_format == "html":
            encoded_css = base64.b64encode(CSS_CONTENT.encode()).decode()
            extra_args.append(f"--css=data:text/css;base64,{encoded_css}")
        pypandoc.convert_file(
            input_path,
            pandoc_output_format,
            outputfile=output_path,
            format=input_format,
            extra_args=extra_args,
        )

    if ext == ".pdf" and os.path.exists(input_path) and input_path.endswith(".docx") and input_path != output_path:
        try:
            os.unlink(input_path)
        except OSError:
            pass


def normalized_extension(path: str) -> str:
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    if ext == "markdown":
        return "md"
    if ext == "htm":
        return "html"
    return ext


def readable_size(path: str) -> str:
    size_bytes = os.path.getsize(path)
    if size_bytes >= 1024 * 1024:
        return f"{size_bytes / (1024 * 1024):.2f} MB"
    return f"{size_bytes / 1024:.2f} KB"


def elide_path(path: str, limit: int = 78) -> str:
    if len(path) <= limit:
        return path
    return "..." + path[-(limit - 3):]


def with_shadow(widget: QWidget, blur: int = 28, y_offset: int = 8):
    shadow = QGraphicsDropShadowEffect(widget)
    shadow.setBlurRadius(blur)
    shadow.setOffset(0, y_offset)
    shadow.setColor(Qt.GlobalColor.transparent)
    widget.setGraphicsEffect(shadow)


class FormatButton(QToolButton):
    def __init__(self, fmt: str):
        super().__init__()
        self.fmt = fmt
        self.meta = FORMAT_META[fmt]
        self.blocked = False
        self.setCheckable(True)
        self.setCursor(Qt.CursorShape.PointingHandCursor)
        self.setToolButtonStyle(Qt.ToolButtonStyle.ToolButtonTextOnly)
        self.setText(f"{self.meta['title']}\n{self.meta['subtitle']}")
        self.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Fixed)
        self.setFixedHeight(64)
        self.toggled.connect(self.refresh_style)
        self.refresh_style()

    def set_blocked(self, blocked: bool):
        self.blocked = blocked
        self.setEnabled(not blocked)
        if blocked:
            self.setChecked(False)
            self.setToolTip("源文件已是该格式")
        else:
            self.setToolTip("")
        self.refresh_style()

    def refresh_style(self):
        color = self.meta["color"]
        if self.blocked:
            self.setStyleSheet(
                """
                QToolButton {
                    background: #F8FAFC;
                    color: #94A3B8;
                    border: 1px solid #E2E8F0;
                    border-radius: 10px;
                    font-family: "Microsoft YaHei UI";
                    font-size: 13px;
                    font-weight: 600;
                    padding: 7px 8px;
                }
                """
            )
        elif self.isChecked():
            self.setStyleSheet(
                f"""
                QToolButton {{
                    background: {color};
                    color: white;
                    border: 1px solid {color};
                    border-radius: 10px;
                    font-family: "Microsoft YaHei UI";
                    font-size: 13px;
                    font-weight: 700;
                    padding: 7px 8px;
                }}
                """
            )
        else:
            self.setStyleSheet(
                f"""
                QToolButton {{
                    background: #FFFFFF;
                    color: #334155;
                    border: 1px solid #CBD5E1;
                    border-radius: 10px;
                    font-family: "Microsoft YaHei UI";
                    font-size: 13px;
                    font-weight: 600;
                    padding: 7px 8px;
                }}
                QToolButton:hover {{
                    border-color: {color};
                    background: #F8FAFC;
                }}
                """
            )


class ConversionWorker(QThread):
    progress_changed = Signal(int)
    conversion_done = Signal(list, str, list)

    def __init__(self, source_file: str, output_dir: str, tasks: list[tuple[str, str]], notices: list[str]):
        super().__init__()
        self.source_file = source_file
        self.output_dir = output_dir
        self.tasks = tasks
        self.notices = notices

    def run(self):
        output_files = []
        errors = list(self.notices)
        total = len(self.tasks)
        completed = 0

        def convert_one(task):
            fmt, output_path = task
            try:
                convert_with_pandoc(self.source_file, fmt, output_path)
                if os.path.exists(output_path):
                    return "success", os.path.basename(output_path)
                return "error", f"{fmt.upper()}: 文件未生成"
            except Exception as exc:
                return "error", f"{fmt.upper()}: {str(exc)[:120]}"

        max_workers = min(3, total) if total else 1
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            future_to_task = {executor.submit(convert_one, task): task for task in self.tasks}
            for future in as_completed(future_to_task):
                status, message = future.result()
                if status == "success":
                    output_files.append(message)
                else:
                    errors.append(message)
                completed += 1
                self.progress_changed.emit(int((completed / total) * 100) if total else 100)

        self.conversion_done.emit(output_files, self.output_dir, errors)


class StyledMessageDialog(QDialog):
    def __init__(self, title: str, message: str, icon_type: str = "info", parent=None):
        super().__init__(parent)
        self.setWindowTitle(title)
        self.setWindowIcon(get_app_icon())
        self.setModal(True)
        self.setFixedWidth(440)
        self.setStyleSheet(APP_QSS)

        # 无边框窗口，隐藏标题栏
        self.setWindowFlags(Qt.WindowType.Dialog | Qt.WindowType.FramelessWindowHint)

        # 添加阴影效果
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(0)

        container = QFrame()
        container.setObjectName("DialogPanel")
        shadow = QGraphicsDropShadowEffect(container)
        shadow.setBlurRadius(40)
        shadow.setOffset(0, 10)
        shadow.setColor(QColor(0, 0, 0, 100))
        container.setGraphicsEffect(shadow)

        panel_layout = QVBoxLayout(container)
        panel_layout.setContentsMargins(0, 0, 0, 0)
        panel_layout.setSpacing(0)
        layout.addWidget(container)

        top_bar = QFrame()
        top_bar.setObjectName("DialogTopBar")
        top_bar.setFixedHeight(5)
        panel_layout.addWidget(top_bar)

        body = QVBoxLayout()
        body.setContentsMargins(30, 26, 30, 26)
        body.setSpacing(0)
        panel_layout.addLayout(body)

        icon_emoji = {"warning": "⚠️", "info": "ℹ️", "error": "❌"}.get(icon_type, "ℹ️")
        icon = QLabel(icon_emoji)
        icon.setObjectName("MessageIcon")
        icon.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(icon)
        body.addSpacing(16)

        title_label = QLabel(title)
        title_label.setObjectName("DialogTitle")
        title_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(title_label)
        body.addSpacing(10)

        message_label = QLabel(message)
        message_label.setObjectName("DialogMessage")
        message_label.setWordWrap(True)
        message_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(message_label)
        body.addSpacing(24)

        ok_btn = QPushButton("确定")
        ok_btn.setObjectName("PrimaryButton")
        ok_btn.setFixedWidth(120)
        ok_btn.clicked.connect(self.accept)
        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        btn_layout.addWidget(ok_btn)
        btn_layout.addStretch()
        body.addLayout(btn_layout)

        # 关闭按钮
        self.close_btn = QPushButton("×", container)
        self.close_btn.setObjectName("DialogCloseButton")
        self.close_btn.setFixedSize(32, 32)
        self.close_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.close_btn.clicked.connect(self.reject)

    def showEvent(self, event):
        super().showEvent(event)
        # 将关闭按钮定位到右上角
        self.close_btn.move(self.close_btn.parent().width() - 40, 8)
        self.close_btn.raise_()
        if self.parent():
            parent_geo = self.parent().geometry()
            x = parent_geo.x() + (parent_geo.width() - self.width()) // 2
            y = parent_geo.y() + (parent_geo.height() - self.height()) // 2
            self.move(x, y)


class OverwriteDialog(QDialog):
    def __init__(self, filename: str, parent=None):
        super().__init__(parent)
        self.choice = "cancel"
        self.setWindowTitle("文件已存在")
        self.setWindowIcon(get_app_icon())
        self.setModal(True)
        self.setFixedWidth(520)
        self.setStyleSheet(APP_QSS)

        # 无边框窗口，隐藏标题栏
        self.setWindowFlags(Qt.WindowType.Dialog | Qt.WindowType.FramelessWindowHint)

        # 添加阴影效果
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(0)

        container = QFrame()
        container.setObjectName("DialogPanel")
        shadow = QGraphicsDropShadowEffect(container)
        shadow.setBlurRadius(40)
        shadow.setOffset(0, 10)
        shadow.setColor(QColor(0, 0, 0, 100))
        container.setGraphicsEffect(shadow)

        panel_layout = QVBoxLayout(container)
        panel_layout.setContentsMargins(0, 0, 0, 0)
        panel_layout.setSpacing(0)
        layout.addWidget(container)

        top_bar = QFrame()
        top_bar.setObjectName("DialogTopBar")
        top_bar.setFixedHeight(5)
        panel_layout.addWidget(top_bar)

        body = QVBoxLayout()
        body.setContentsMargins(30, 24, 30, 26)
        body.setSpacing(0)
        panel_layout.addLayout(body)

        icon = QLabel("⚠️")
        icon.setObjectName("WarningBadge")
        icon.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(icon)
        body.addSpacing(14)

        title = QLabel("文件已存在")
        title.setObjectName("DialogTitle")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(title)
        body.addSpacing(8)

        desc = QLabel("目标目录中已有同名文件，请选择处理方式：")
        desc.setObjectName("DialogSubtitle")
        desc.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(desc)
        body.addSpacing(20)

        file_box = QFrame()
        file_box.setObjectName("NoticePanel")
        file_inner = QHBoxLayout(file_box)
        file_inner.setContentsMargins(16, 14, 16, 14)
        file_inner.setSpacing(12)

        file_icon = QLabel("📄")
        file_icon.setObjectName("FileIconLarge")
        file_icon.setAlignment(Qt.AlignmentFlag.AlignCenter)
        file_inner.addWidget(file_icon)

        file_info = QVBoxLayout()
        file_info.setSpacing(4)
        file_label = QLabel(filename)
        file_label.setObjectName("DialogFileName")
        file_label.setWordWrap(True)
        file_detail = QLabel("将被新生成的文件影响")
        file_detail.setObjectName("DialogFileDetail")
        file_info.addWidget(file_label)
        file_info.addWidget(file_detail)
        file_inner.addLayout(file_info, stretch=1)
        body.addWidget(file_box)
        body.addSpacing(24)

        separator = QFrame()
        separator.setObjectName("DialogSeparator")
        separator.setFixedHeight(1)
        body.addWidget(separator)
        body.addSpacing(20)

        actions = QHBoxLayout()
        actions.setSpacing(10)

        cancel_btn = QPushButton("取消")
        cancel_btn.setObjectName("GhostButton")
        cancel_btn.setFixedHeight(40)
        cancel_btn.clicked.connect(lambda: self.finish("cancel"))
        actions.addWidget(cancel_btn)

        actions.addStretch()

        skip_btn = QPushButton("跳过")
        skip_btn.setObjectName("SecondaryButton")
        skip_btn.setFixedHeight(40)
        skip_btn.clicked.connect(lambda: self.finish("skip"))
        actions.addWidget(skip_btn)

        overwrite_btn = QPushButton("覆盖替换")
        overwrite_btn.setObjectName("DangerButton")
        overwrite_btn.setFixedHeight(40)
        overwrite_btn.clicked.connect(lambda: self.finish("overwrite"))
        actions.addWidget(overwrite_btn)

        rename_btn = QPushButton("智能重命名")
        rename_btn.setObjectName("PrimaryButton")
        rename_btn.setFixedHeight(40)
        rename_btn.clicked.connect(lambda: self.finish("rename"))
        actions.addWidget(rename_btn)

        body.addLayout(actions)

        # 关闭按钮
        self.close_btn = QPushButton("×", container)
        self.close_btn.setObjectName("DialogCloseButton")
        self.close_btn.setFixedSize(32, 32)
        self.close_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.close_btn.clicked.connect(lambda: self.finish("cancel"))

    def showEvent(self, event):
        super().showEvent(event)
        # 将关闭按钮定位到右上角
        self.close_btn.move(self.close_btn.parent().width() - 40, 8)
        self.close_btn.raise_()
        if self.parent():
            parent_geo = self.parent().geometry()
            x = parent_geo.x() + (parent_geo.width() - self.width()) // 2
            y = parent_geo.y() + (parent_geo.height() - self.height()) // 2
            self.move(x, y)

    def finish(self, choice: str):
        self.choice = choice
        self.accept()


class ResultDialog(QDialog):
    def __init__(self, output_files: list[str], output_dir: str, notices: list[str], parent=None):
        super().__init__(parent)
        self.output_dir = output_dir
        self.setWindowTitle("转换完成")
        self.setWindowIcon(get_app_icon())
        self.setModal(True)
        self.setFixedWidth(540)
        self.setStyleSheet(APP_QSS)

        # 无边框窗口，隐藏标题栏
        self.setWindowFlags(Qt.WindowType.Dialog | Qt.WindowType.FramelessWindowHint)

        # 添加阴影效果
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(0)

        container = QFrame()
        container.setObjectName("DialogPanel")
        shadow = QGraphicsDropShadowEffect(container)
        shadow.setBlurRadius(40)
        shadow.setOffset(0, 10)
        shadow.setColor(QColor(0, 0, 0, 100))
        container.setGraphicsEffect(shadow)

        panel_layout = QVBoxLayout(container)
        panel_layout.setContentsMargins(0, 0, 0, 0)
        panel_layout.setSpacing(0)
        layout.addWidget(container)

        top_bar = QFrame()
        top_bar.setObjectName("DialogTopBar")
        top_bar.setFixedHeight(5)
        panel_layout.addWidget(top_bar)

        body = QVBoxLayout()
        body.setContentsMargins(30, 26, 30, 26)
        body.setSpacing(0)
        panel_layout.addLayout(body)

        icon = QLabel("✅")
        icon.setObjectName("SuccessIcon")
        icon.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(icon)
        body.addSpacing(14)

        title = QLabel("转换完成")
        title.setObjectName("DialogTitle")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(title)
        body.addSpacing(18)

        if output_files:
            file_title = QLabel(f"成功生成 {len(output_files)} 个文件")
            file_title.setObjectName("SectionTitle")
            body.addWidget(file_title)
            body.addSpacing(10)

            file_list = QFrame()
            file_list.setObjectName("ListPanel")
            file_layout = QVBoxLayout(file_list)
            file_layout.setContentsMargins(14, 12, 14, 12)
            file_layout.setSpacing(8)
            for file_name in output_files:
                row = QLabel(file_name)
                row.setObjectName("SuccessRow")
                row.setMinimumHeight(34)
                file_layout.addWidget(row)
            body.addWidget(file_list)
            body.addSpacing(14)

        dir_label = QLabel(f"📁 {elide_path(output_dir, 70)}")
        dir_label.setObjectName("PathText")
        dir_label.setWordWrap(True)
        body.addWidget(dir_label)

        if notices:
            body.addSpacing(18)
            notice_title = QLabel("未完成项目")
            notice_title.setObjectName("SectionTitle")
            body.addWidget(notice_title)
            body.addSpacing(8)
            for item in notices[:5]:
                warning = QLabel(item)
                warning.setObjectName("WarningText")
                warning.setWordWrap(True)
                body.addWidget(warning)

        body.addSpacing(22)

        actions = QHBoxLayout()
        actions.setSpacing(10)
        actions.addStretch()

        open_btn = QPushButton("打开文件夹")
        open_btn.setObjectName("PrimaryButton")
        open_btn.setFixedHeight(40)
        open_btn.clicked.connect(self.open_folder)
        actions.addWidget(open_btn)

        ok_btn = QPushButton("确定")
        ok_btn.setObjectName("SecondaryButton")
        ok_btn.setFixedHeight(40)
        ok_btn.clicked.connect(self.accept)
        actions.addWidget(ok_btn)
        body.addLayout(actions)

        # 关闭按钮
        self.close_btn = QPushButton("×", container)
        self.close_btn.setObjectName("DialogCloseButton")
        self.close_btn.setFixedSize(32, 32)
        self.close_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.close_btn.clicked.connect(self.accept)

    def showEvent(self, event):
        super().showEvent(event)
        # 将关闭按钮定位到右上角
        self.close_btn.move(self.close_btn.parent().width() - 40, 8)
        self.close_btn.raise_()
        if self.parent():
            parent_geo = self.parent().geometry()
            self.move(
                parent_geo.x() + (parent_geo.width() - self.width()) // 2,
                parent_geo.y() + (parent_geo.height() - self.height()) // 2
            )

    def open_folder(self):
        QDesktopServices.openUrl(QUrl.fromLocalFile(self.output_dir))
        self.accept()


class ConverterApp(QMainWindow):
    def __init__(self):
        super().__init__()
        self.file_path = None
        self.custom_output_dir = None
        self.worker = None

        self.setWindowTitle(f"文件格式转换器 v{get_current_version()}")
        self.setWindowIcon(get_app_icon())
        self.setMinimumSize(1080, 720)
        self.resize(1120, 760)
        self.setAcceptDrops(True)
        self.setStyleSheet(APP_QSS)

        self.format_buttons = {}
        self.setup_ui()
        self.update_output_label()

        # 启动时检查更新（在后台线程）
        if UPDATER_AVAILABLE:
            self.check_updates_on_startup()

    def setup_ui(self):
        root = QWidget()
        root.setObjectName("AppRoot")
        self.setCentralWidget(root)

        page = QVBoxLayout(root)
        page.setContentsMargins(34, 28, 34, 28)
        page.setSpacing(22)

        header = QHBoxLayout()
        header.setSpacing(16)

        title_box = QVBoxLayout()
        title_box.setSpacing(5)
        title = QLabel("文件格式转换器")
        title.setObjectName("AppTitle")
        subtitle = QLabel("Markdown / HTML / Word / PDF / TXT")
        subtitle.setObjectName("AppSubtitle")
        title_box.addWidget(title)
        title_box.addWidget(subtitle)
        header.addLayout(title_box)
        header.addStretch()

        # 帮助按钮
        help_btn = QPushButton("帮助")
        help_btn.setObjectName("HelpButton")
        help_btn.setFixedSize(68, 34)
        help_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        help_menu = QMenu(help_btn)
        help_menu.setStyleSheet("""
            QMenu {
                background: #FFFFFF;
                border: 1px solid #E2E8F0;
                border-radius: 8px;
                padding: 8px;
            }
            QMenu::item {
                padding: 8px 20px;
                border-radius: 6px;
                color: #334155;
            }
            QMenu::item:selected {
                background: #F1F5F9;
            }
        """)

        about_action = QAction("关于", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)

        if UPDATER_AVAILABLE:
            check_update_action = QAction("检查更新", self)
            check_update_action.triggered.connect(self.manual_check_updates)
            help_menu.addAction(check_update_action)

        help_btn.setMenu(help_menu)
        header.addWidget(help_btn)

        header.addSpacing(8)

        self.status_chip = QLabel("就绪")
        self.status_chip.setObjectName("StatusChip")
        self.status_chip.setFixedSize(88, 34)
        self.status_chip.setAlignment(Qt.AlignmentFlag.AlignCenter)
        header.addWidget(self.status_chip)
        page.addLayout(header)

        content = QHBoxLayout()
        content.setSpacing(18)
        page.addLayout(content, stretch=1)

        main_card = self.make_card()
        main_layout = QVBoxLayout(main_card)
        main_layout.setContentsMargins(24, 22, 24, 22)
        main_layout.setSpacing(18)
        content.addWidget(main_card, stretch=5)

        file_section = self.make_section("源文件")
        file_layout = QVBoxLayout(self.section_body(file_section))
        file_layout.setContentsMargins(16, 14, 16, 14)
        file_layout.setSpacing(12)

        choose_row = QHBoxLayout()
        choose_row.setSpacing(12)
        self.file_name = QLabel("未选择文件")
        self.file_name.setObjectName("FileName")
        self.file_name.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        choose_row.addWidget(self.file_name, stretch=1)

        self.choose_btn = QPushButton("选择文件")
        self.choose_btn.setObjectName("PrimaryButton")
        self.choose_btn.clicked.connect(self.select_file)
        choose_row.addWidget(self.choose_btn)
        file_layout.addLayout(choose_row)

        self.file_meta = QLabel("类型：-    大小：-")
        self.file_meta.setObjectName("MetaText")
        file_layout.addWidget(self.file_meta)

        self.file_path_label = QLabel("路径：-")
        self.file_path_label.setObjectName("PathText")
        self.file_path_label.setWordWrap(True)
        self.file_path_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        file_layout.addWidget(self.file_path_label)
        main_layout.addWidget(file_section)

        format_section = self.make_section("输出格式")
        format_body = self.section_body(format_section)
        format_body.setFixedHeight(96)
        format_layout = QGridLayout(format_body)
        format_layout.setContentsMargins(14, 14, 14, 14)
        format_layout.setHorizontalSpacing(12)
        format_layout.setVerticalSpacing(0)

        for idx, fmt in enumerate(OUTPUT_FORMATS):
            button = FormatButton(fmt)
            self.format_buttons[fmt] = button
            format_layout.addWidget(button, 0, idx)
        main_layout.addWidget(format_section)

        output_section = self.make_section("保存位置")
        output_layout = QHBoxLayout(self.section_body(output_section))
        output_layout.setContentsMargins(16, 14, 16, 14)
        output_layout.setSpacing(12)

        self.output_label = QLabel("")
        self.output_label.setObjectName("PathText")
        self.output_label.setMinimumHeight(36)
        self.output_label.setWordWrap(True)
        self.output_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextSelectableByMouse)
        output_layout.addWidget(self.output_label, stretch=1)

        output_btn = QPushButton("更改")
        output_btn.setObjectName("SecondaryButton")
        output_btn.clicked.connect(self.select_output_dir)
        output_layout.addWidget(output_btn)
        main_layout.addWidget(output_section)

        self.progress = QProgressBar()
        self.progress.setObjectName("MainProgress")
        self.progress.setRange(0, 100)
        self.progress.setValue(0)
        self.progress.setTextVisible(False)
        main_layout.addWidget(self.progress)

        bottom = QHBoxLayout()
        bottom.setSpacing(12)
        self.status_text = QLabel("就绪")
        self.status_text.setObjectName("MetaText")
        bottom.addWidget(self.status_text, stretch=1)

        self.convert_btn = QPushButton("开始转换")
        self.convert_btn.setObjectName("ConvertButton")
        self.convert_btn.clicked.connect(self.convert)
        bottom.addWidget(self.convert_btn)
        main_layout.addLayout(bottom)

        side_panel = self.make_card()
        side_panel.setObjectName("SidePanel")
        side_layout = QVBoxLayout(side_panel)
        side_layout.setContentsMargins(22, 22, 22, 22)
        side_layout.setSpacing(14)
        content.addWidget(side_panel, stretch=2)

        summary_title = QLabel("当前任务")
        summary_title.setObjectName("PanelTitle")
        side_layout.addWidget(summary_title)

        self.summary_file = QLabel("文件：-")
        self.summary_file.setObjectName("PanelText")
        self.summary_file.setWordWrap(True)
        side_layout.addWidget(self.summary_file)

        self.summary_formats = QLabel("格式：-")
        self.summary_formats.setObjectName("PanelText")
        self.summary_formats.setWordWrap(True)
        side_layout.addWidget(self.summary_formats)

        self.summary_output = QLabel("位置：-")
        self.summary_output.setObjectName("PanelText")
        self.summary_output.setWordWrap(True)
        side_layout.addWidget(self.summary_output)
        side_layout.addStretch()

        hint = QLabel("可将文件拖入窗口")
        hint.setObjectName("DropHint")
        hint.setFixedHeight(116)
        hint.setAlignment(Qt.AlignmentFlag.AlignCenter)
        side_layout.addWidget(hint)

        for button in self.format_buttons.values():
            button.toggled.connect(self.update_summary)

    # 移除 setup_menu，改用按钮形式

    def make_card(self) -> QFrame:
        card = QFrame()
        card.setObjectName("Card")
        card.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        return card

    def make_section(self, title: str) -> QFrame:
        section = QFrame()
        section.setObjectName("Section")
        layout = QVBoxLayout(section)
        layout.setContentsMargins(0, 0, 0, 0)
        layout.setSpacing(0)

        heading = QLabel(title)
        heading.setObjectName("SectionTitle")
        heading.setContentsMargins(0, 0, 0, 8)
        layout.addWidget(heading)

        body = QFrame()
        body.setObjectName("SectionBody")
        layout.addWidget(body)
        return section

    def section_body(self, section: QFrame) -> QFrame:
        return section.findChild(QFrame, "SectionBody")

    def select_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择要转换的文件", "", FILE_DIALOG_FILTER)
        if file_path:
            self.set_file(file_path)

    def set_file(self, file_path: str):
        self.file_path = file_path
        file_name = os.path.basename(file_path)
        file_ext = os.path.splitext(file_path)[1].upper() or "-"

        self.file_name.setText(file_name)
        self.file_meta.setText(f"类型：{file_ext}    大小：{readable_size(file_path)}")
        self.file_path_label.setText(f"路径：{file_path}")
        self.progress.setValue(0)
        self.set_status("就绪", "ready")

        if not self.custom_output_dir:
            self.update_output_label(os.path.dirname(file_path))
        else:
            self.update_output_label(self.custom_output_dir)

        source_ext = normalized_extension(file_path)
        for fmt, button in self.format_buttons.items():
            button.set_blocked(fmt == source_ext)

        self.update_summary()

    def select_output_dir(self):
        directory = QFileDialog.getExistingDirectory(self, "选择保存位置", self.current_output_dir())
        if directory:
            self.custom_output_dir = directory
            self.update_output_label(directory)
            self.update_summary()

    def update_output_label(self, directory: str | None = None):
        if directory is None:
            directory = self.current_output_dir()
        self.output_label.setText(directory if directory else "选择文件后使用源文件所在位置")

    def current_output_dir(self) -> str:
        if self.custom_output_dir:
            return self.custom_output_dir
        if self.file_path:
            return os.path.dirname(self.file_path)
        return ""

    def selected_formats(self) -> list[str]:
        return [fmt for fmt, button in self.format_buttons.items() if button.isChecked() and button.isEnabled()]

    def update_summary(self):
        file_text = os.path.basename(self.file_path) if self.file_path else "-"
        formats = self.selected_formats()
        output_dir = self.current_output_dir() or "-"

        self.summary_file.setText(f"文件：{file_text}")
        self.summary_formats.setText(f"格式：{', '.join(fmt.upper() for fmt in formats) if formats else '-'}")
        self.summary_output.setText(f"位置：{elide_path(output_dir, 64) if output_dir != '-' else '-'}")

    def convert(self):
        if self.worker and self.worker.isRunning():
            return

        if not self.file_path:
            StyledMessageDialog("提示", "请先选择要转换的文件。", "warning", self).exec()
            return

        selected = self.selected_formats()
        if not selected:
            StyledMessageDialog("提示", "请至少选择一种输出格式。", "warning", self).exec()
            return

        source_ext = normalized_extension(self.file_path)
        formats_to_convert = [fmt for fmt in selected if fmt != source_ext]
        if not formats_to_convert:
            StyledMessageDialog("提示", "输出格式与源文件格式相同，没有生成新文件。", "info", self).exec()
            return

        output_dir = self.current_output_dir()
        if not output_dir:
            StyledMessageDialog("提示", "请选择保存位置。", "warning", self).exec()
            return

        base_name = os.path.splitext(os.path.basename(self.file_path))[0]
        tasks = []
        notices = []

        for fmt in formats_to_convert:
            output_path = os.path.join(output_dir, f"{base_name}.{fmt}")
            if os.path.exists(output_path):
                choice = self.ask_overwrite(os.path.basename(output_path))
                if choice == "cancel":
                    self.set_status("已取消", "idle")
                    return
                if choice == "skip":
                    notices.append(f"{fmt.upper()}: 已跳过同名文件")
                    continue
                if choice == "rename":
                    output_path = self.available_output_path(output_dir, base_name, fmt)
            tasks.append((fmt, output_path))

        if not tasks:
            StyledMessageDialog("提示", "没有需要转换的输出格式。", "info", self).exec()
            self.set_status("就绪", "ready")
            return

        self.set_busy(True)
        self.worker = ConversionWorker(self.file_path, output_dir, tasks, notices)
        self.worker.progress_changed.connect(self.progress.setValue)
        self.worker.conversion_done.connect(self.on_conversion_done)
        self.worker.start()

    def ask_overwrite(self, filename: str) -> str:
        dialog = OverwriteDialog(filename, self)
        dialog.exec()
        return dialog.choice

    def available_output_path(self, output_dir: str, base_name: str, fmt: str) -> str:
        counter = 1
        while True:
            candidate = os.path.join(output_dir, f"{base_name}_{counter}.{fmt}")
            if not os.path.exists(candidate):
                return candidate
            counter += 1

    def set_busy(self, busy: bool):
        self.choose_btn.setEnabled(not busy)
        self.convert_btn.setEnabled(not busy)
        for button in self.format_buttons.values():
            button.setEnabled((not busy) and (not button.blocked))
            button.refresh_style()

        if busy:
            self.convert_btn.setText("转换中...")
            self.set_status("正在转换", "working")
            self.progress.setValue(0)
        else:
            self.convert_btn.setText("开始转换")

    def set_status(self, text: str, state: str):
        self.status_text.setText(text)
        self.status_chip.setText(text)
        self.status_chip.setProperty("state", state)
        self.status_chip.style().unpolish(self.status_chip)
        self.status_chip.style().polish(self.status_chip)

    def on_conversion_done(self, output_files: list[str], output_dir: str, notices: list[str]):
        self.set_busy(False)

        if output_files:
            self.progress.setValue(100)
            self.set_status("转换成功", "success")
            ResultDialog(output_files, output_dir, notices, self).exec()
        elif notices:
            self.progress.setValue(0)
            self.set_status("未生成文件", "warning")
            StyledMessageDialog("转换未完成", "\n".join(notices), "warning", self).exec()
        else:
            self.progress.setValue(0)
            self.set_status("转换失败", "error")
            StyledMessageDialog("转换失败", "所有格式转换失败。", "error", self).exec()

    def dragEnterEvent(self, event: QDragEnterEvent):
        if event.mimeData().hasUrls():
            for url in event.mimeData().urls():
                path = url.toLocalFile()
                if os.path.splitext(path)[1].lower() in SUPPORTED_EXTENSIONS:
                    event.acceptProposedAction()
                    return
        event.ignore()

    def dropEvent(self, event: QDropEvent):
        for url in event.mimeData().urls():
            path = url.toLocalFile()
            if os.path.isfile(path) and os.path.splitext(path)[1].lower() in SUPPORTED_EXTENSIONS:
                self.set_file(path)
                event.acceptProposedAction()
                return
        event.ignore()

    def check_updates_on_startup(self):
        """启动时在后台检查更新"""
        from threading import Thread

        def check():
            has_update, latest, download_url, notes = check_for_updates()
            if has_update:
                # 在主线程显示更新对话框
                QApplication.instance().postEvent(
                    self,
                    UpdateAvailableEvent(latest, download_url, notes)
                )

        Thread(target=check, daemon=True).start()

    def manual_check_updates(self):
        """手动检查更新"""
        if not UPDATER_AVAILABLE:
            return

        has_update, latest, download_url, notes = check_for_updates(timeout=10)

        if has_update:
            self.show_update_dialog(latest, download_url, notes)
        else:
            StyledMessageDialog(
                "已是最新版本",
                f"当前版本 {get_current_version()} 已是最新版本。",
                "info",
                self
            ).exec()

    def show_update_dialog(self, version, download_url, notes):
        """显示更新对话框"""
        dialog = UpdateDialog(version, download_url, notes, self)
        if dialog.exec() == QDialog.DialogCode.Accepted:
            open_download_page(download_url)

    def show_about(self):
        """显示关于对话框"""
        about_text = (
            f"文件格式转换器 v{get_current_version()}\n\n"
            "支持 Markdown、HTML、Word、PDF、TXT 等格式互转\n\n"
            "本软件仅供学习交流使用\n"
            "使用本软件产生的任何后果由使用者自行承担"
        )
        StyledMessageDialog("关于", about_text, "info", self).exec()

    def customEvent(self, event):
        """处理自定义事件（更新通知）"""
        if isinstance(event, UpdateAvailableEvent):
            self.show_update_dialog(event.version, event.download_url, event.notes)


class UpdateAvailableEvent:
    """自定义事件：有可用更新"""
    def __init__(self, version, download_url, notes):
        self.version = version
        self.download_url = download_url
        self.notes = notes
        self.type_id = QThread.eventType()


class UpdateDialog(QDialog):
    """更新提示对话框"""
    def __init__(self, version, download_url, notes, parent=None):
        super().__init__(parent)
        self.version = version
        self.download_url = download_url
        self.notes = notes

        self.setWindowTitle("发现新版本")
        self.setWindowIcon(get_app_icon())
        self.setModal(True)
        self.setFixedWidth(520)
        self.setStyleSheet(APP_QSS)
        self.setWindowFlags(Qt.WindowType.Dialog | Qt.WindowType.FramelessWindowHint)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)

        layout = QVBoxLayout(self)
        layout.setContentsMargins(20, 20, 20, 20)
        layout.setSpacing(0)

        container = QFrame()
        container.setObjectName("DialogPanel")
        shadow = QGraphicsDropShadowEffect(container)
        shadow.setBlurRadius(40)
        shadow.setOffset(0, 10)
        shadow.setColor(QColor(0, 0, 0, 100))
        container.setGraphicsEffect(shadow)

        panel_layout = QVBoxLayout(container)
        panel_layout.setContentsMargins(0, 0, 0, 0)
        panel_layout.setSpacing(0)
        layout.addWidget(container)

        top_bar = QFrame()
        top_bar.setObjectName("DialogTopBar")
        top_bar.setFixedHeight(5)
        panel_layout.addWidget(top_bar)

        body = QVBoxLayout()
        body.setContentsMargins(30, 24, 30, 26)
        body.setSpacing(0)
        panel_layout.addLayout(body)

        icon = QLabel("🎉")
        icon.setObjectName("MessageIcon")
        icon.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(icon)
        body.addSpacing(14)

        title = QLabel("发现新版本")
        title.setObjectName("DialogTitle")
        title.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(title)
        body.addSpacing(8)

        version_label = QLabel(f"v{version} 已发布")
        version_label.setObjectName("DialogSubtitle")
        version_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        body.addWidget(version_label)
        body.addSpacing(20)

        if notes and notes.strip():
            notes_box = QFrame()
            notes_box.setObjectName("ListPanel")
            notes_layout = QVBoxLayout(notes_box)
            notes_layout.setContentsMargins(16, 14, 16, 14)

            notes_title = QLabel("更新内容:")
            notes_title.setObjectName("SectionTitle")
            notes_layout.addWidget(notes_title)
            notes_layout.addSpacing(8)

            # 使用滚动区域防止内容过长
            from PySide6.QtWidgets import QScrollArea, QTextEdit
            notes_text = QTextEdit()
            notes_text.setObjectName("DialogMessage")
            notes_text.setReadOnly(True)
            notes_text.setPlainText(notes[:800])  # 限制长度
            notes_text.setMaximumHeight(150)
            notes_text.setMinimumHeight(100)
            notes_text.setFrameStyle(0)
            notes_text.setStyleSheet("""
                QTextEdit {
                    background: transparent;
                    border: none;
                    color: #475569;
                    font-size: 13px;
                    line-height: 1.6;
                }
                QScrollBar:vertical {
                    background: #F1F5F9;
                    width: 8px;
                    border-radius: 4px;
                    margin: 0px;
                }
                QScrollBar::handle:vertical {
                    background: #CBD5E1;
                    border-radius: 4px;
                    min-height: 20px;
                }
                QScrollBar::handle:vertical:hover {
                    background: #94A3B8;
                }
                QScrollBar::add-line:vertical,
                QScrollBar::sub-line:vertical {
                    height: 0px;
                }
                QScrollBar::add-page:vertical,
                QScrollBar::sub-page:vertical {
                    background: transparent;
                }
            """)
            notes_layout.addWidget(notes_text)

            body.addWidget(notes_box)
            body.addSpacing(20)

        actions = QHBoxLayout()
        actions.setSpacing(10)

        later_btn = QPushButton("稍后提醒")
        later_btn.setObjectName("SecondaryButton")
        later_btn.setFixedHeight(40)
        later_btn.clicked.connect(self.reject)
        actions.addWidget(later_btn)

        download_btn = QPushButton("立即下载")
        download_btn.setObjectName("PrimaryButton")
        download_btn.setFixedHeight(40)
        download_btn.clicked.connect(self.accept)
        actions.addWidget(download_btn)

        body.addLayout(actions)

        self.close_btn = QPushButton("×", container)
        self.close_btn.setObjectName("DialogCloseButton")
        self.close_btn.setFixedSize(32, 32)
        self.close_btn.setCursor(Qt.CursorShape.PointingHandCursor)
        self.close_btn.clicked.connect(self.reject)

    def showEvent(self, event):
        super().showEvent(event)
        self.close_btn.move(self.close_btn.parent().width() - 40, 8)
        self.close_btn.raise_()
        if self.parent():
            parent_geo = self.parent().geometry()
            x = parent_geo.x() + (parent_geo.width() - self.width()) // 2
            y = parent_geo.y() + (parent_geo.height() - self.height()) // 2
            self.move(x, y)


APP_QSS = """
QWidget#AppRoot {
    background: #F4F7FB;
    color: #111827;
    font-family: "Microsoft YaHei UI";
    font-size: 13px;
}
QLabel#AppTitle {
    color: #0F172A;
    font-size: 26px;
    font-weight: 800;
}
QLabel#AppSubtitle {
    color: #64748B;
    font-size: 13px;
}
QLabel#StatusChip {
    min-width: 86px;
    min-height: 30px;
    border-radius: 15px;
    background: #E2E8F0;
    color: #475569;
    font-weight: 700;
}
QLabel#StatusChip[state="working"] {
    background: #FEF3C7;
    color: #92400E;
}
QLabel#StatusChip[state="success"] {
    background: #DCFCE7;
    color: #166534;
}
QLabel#StatusChip[state="error"] {
    background: #FEE2E2;
    color: #991B1B;
}
QLabel#StatusChip[state="warning"] {
    background: #FFEDD5;
    color: #9A3412;
}
QPushButton#HelpButton {
    background: #2563EB;
    color: white;
    border: 1px solid #2563EB;
    border-radius: 8px;
    font-weight: 700;
    padding: 0 16px;
    text-align: center;
}
QPushButton#HelpButton:hover {
    background: #1D4ED8;
    border-color: #1D4ED8;
}
QPushButton#HelpButton::menu-indicator {
    image: none;
    width: 0px;
}
QFrame#Card {
    background: #FFFFFF;
    border: 1px solid #E2E8F0;
    border-radius: 12px;
}
QFrame#SidePanel {
    background: #FFFFFF;
    border: 1px solid #E2E8F0;
    border-radius: 12px;
}
QFrame#Section {
    background: transparent;
    border: none;
}
QFrame#SectionBody {
    background: #F8FAFC;
    border: 1px solid #E2E8F0;
    border-radius: 10px;
}
QLabel#SectionTitle {
    color: #0F172A;
    font-size: 14px;
    font-weight: 800;
}
QLabel#FileName {
    color: #0F172A;
    font-size: 16px;
    font-weight: 800;
}
QLabel#MetaText,
QLabel#PathText,
QLabel#PanelText {
    color: #64748B;
    font-size: 12px;
}
QLabel#PanelTitle {
    color: #0F172A;
    font-size: 17px;
    font-weight: 800;
}
QLabel#DropHint {
    min-height: 92px;
    color: #475569;
    background: #F8FAFC;
    border: 1px dashed #94A3B8;
    border-radius: 10px;
    font-weight: 700;
}
QPushButton {
    min-height: 34px;
    border-radius: 8px;
    padding: 0 16px;
    font-weight: 700;
}
QPushButton#PrimaryButton,
QPushButton#ConvertButton {
    background: #2563EB;
    color: white;
    border: 1px solid #2563EB;
}
QPushButton#PrimaryButton:hover,
QPushButton#ConvertButton:hover {
    background: #1D4ED8;
    border-color: #1D4ED8;
}
QPushButton#PrimaryButton:disabled,
QPushButton#ConvertButton:disabled {
    background: #93C5FD;
    border-color: #93C5FD;
}
QPushButton#ConvertButton {
    min-width: 180px;
    min-height: 42px;
    font-size: 14px;
}
QPushButton#SecondaryButton {
    background: #FFFFFF;
    color: #334155;
    border: 1px solid #CBD5E1;
}
QPushButton#SecondaryButton:hover {
    background: #F8FAFC;
    border-color: #94A3B8;
}
QProgressBar#MainProgress {
    height: 10px;
    background: #E2E8F0;
    border: none;
    border-radius: 5px;
}
QProgressBar#MainProgress::chunk {
    background: #2563EB;
    border-radius: 5px;
}
QFrame#DialogPanel {
    background: #FAFBFC;
    border: none;
    border-radius: 12px;
}
QFrame#DialogTopBar {
    background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
        stop:0 #2563EB, stop:0.5 #7C3AED, stop:1 #EC4899);
    border: none;
    border-top-left-radius: 12px;
    border-top-right-radius: 12px;
}
QLabel#DialogTitle {
    color: #0F172A;
    font-size: 20px;
    font-weight: 800;
}
QLabel#DialogSubtitle {
    color: #64748B;
    font-size: 13px;
    line-height: 1.5;
}
QLabel#DialogMessage {
    color: #475569;
    font-size: 13px;
    line-height: 1.6;
}
QLabel#MessageIcon {
    font-size: 48px;
    min-height: 56px;
}
QLabel#WarningBadge {
    font-size: 48px;
    min-height: 56px;
}
QLabel#SuccessIcon {
    font-size: 48px;
    min-height: 56px;
}
QFrame#IconContainer {
    background: transparent;
    border: none;
}
QFrame#NoticePanel {
    background: qlineargradient(x1:0, y1:0, x2:0, y2:1,
        stop:0 #FFFBEB, stop:1 #FEF3C7);
    border: 2px solid #FCD34D;
    border-radius: 12px;
}
QLabel#FileIconLarge {
    font-size: 28px;
    min-width: 44px;
    max-width: 44px;
    min-height: 44px;
    max-height: 44px;
}
QLabel#FileIcon {
    font-size: 22px;
    background: #EFF6FF;
    border-radius: 20px;
    border: 1px solid #BFDBFE;
}
QLabel#DialogFileName {
    color: #0F172A;
    font-size: 14px;
    font-weight: 700;
}
QLabel#DialogFileDetail {
    color: #92400E;
    font-size: 12px;
}
QFrame#DialogSeparator {
    background: #D1D5DB;
    border: none;
}
QPushButton#DangerButton {
    background: #FEE2E2;
    color: #991B1B;
    border: 2px solid #F87171;
    border-radius: 8px;
    padding: 0 20px;
    font-weight: 700;
}
QPushButton#DangerButton:hover {
    background: #FEF2F2;
    border-color: #DC2626;
    color: #7F1D1D;
}
QPushButton#GhostButton {
    background: transparent;
    color: #64748B;
    border: 1px solid transparent;
    border-radius: 8px;
    padding: 0 20px;
    font-weight: 700;
}
QPushButton#GhostButton:hover {
    background: #F1F5F9;
    color: #334155;
}
QPushButton#DialogCloseButton {
    background: transparent;
    color: #94A3B8;
    border: none;
    border-radius: 16px;
    font-size: 24px;
    font-weight: 700;
    padding: 0;
}
QPushButton#DialogCloseButton:hover {
    background: #FEE2E2;
    color: #DC2626;
}
QFrame#ListPanel {
    background: #F8FAFC;
    border: 1px solid #E2E8F0;
    border-radius: 10px;
}
QLabel#SuccessRow {
    color: #166534;
    background: #DCFCE7;
    border-radius: 7px;
    padding-left: 12px;
    font-weight: 700;
}
QLabel#WarningText {
    color: #B45309;
    font-size: 12px;
}
QMessageBox {
    background: #FFFFFF;
    font-family: "Microsoft YaHei UI";
}
QMessageBox QLabel {
    color: #111827;
    font-family: "Microsoft YaHei UI";
}
QDialog {
    background: #FFFFFF;
    font-family: "Microsoft YaHei UI";
}
"""


def main():
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    app.setWindowIcon(get_app_icon())
    window = ConverterApp()
    window.show()
    sys.exit(app.exec())


if __name__ == "__main__":
    main()
